# generators/generateur_word.py
"""
Générateur Word — description technique et fonctionnelle de l'offre Orange Business
Document modifiable, 1 configuration unique.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xFF, 0x79, 0x00)   # Orange officiel
DARK   = RGBColor(0x1A, 0x1A, 0x1A)   # Quasi-noir Orange
BLANC  = RGBColor(0xFF, 0xFF, 0xFF)
SAUMON = RGBColor(0xFF, 0xCC, 0xAA)   # Orange clair (lignes alternées)


def generer_word(data: dict, textes: dict, chemin: str):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.8)
        sec.left_margin = sec.right_margin = Cm(2)

    # ── En-tête ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    _run(p, "Orange Business", size=18, bold=True, color=ORANGE)
    _run(p, "  ·  Tunisie  |  Description Technique et Fonctionnelle", size=10, color=DARK)

    p2 = doc.add_paragraph()
    _run(p2, "Client : ", size=9, bold=True, color=DARK)
    _run(p2, f"{data['client']}   |   Date : {data['date']}   |   Validité : {data['validite']}", size=9, color=DARK)
    _hr(doc)

    # ── Profil client ─────────────────────────────────────────────
    _h1(doc, "PROFIL CLIENT")
    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    for j, (h, v) in enumerate(zip(
        ["Entreprise", "Secteur", "Taille", "Urgence"],
        [data["client"], data["secteur"], data["taille"], (data["urgence"] or "").capitalize()]
    )):
        _cell(t, 0, j, h, bold=True, bg="FF7900", color=BLANC)
        _cell(t, 1, j, v)
    doc.add_paragraph()

    # ── Description technique Fibre ───────────────────────────────
    fibre = data.get("fibre")
    if fibre:
        _h1(doc, "SOLUTION FIBRE OPTIQUE FTTO — DESCRIPTION TECHNIQUE")
        tf = doc.add_table(rows=2, cols=5)
        tf.style = "Table Grid"
        for j, (h, v) in enumerate(zip(
            ["Offre", "Débit", "Engagement", "Distance", "Prix annuel (TND)"],
            [
                fibre.get("nom_offre", "—"),
                f"{fibre.get('debit_mbps', '')} Mbps",
                f"{fibre.get('engagement_mois', '')} mois",
                f"{fibre.get('distance_metres', '')} m",
                f"{fibre.get('prix_annuel', 0):,.0f}",
            ]
        )):
            _cell(tf, 0, j, h, bold=True, bg="FF7900", color=BLANC)
            _cell(tf, 1, j, v)
        doc.add_paragraph()

        _h2(doc, "Caractéristiques techniques")
        details = [
            f"Technologie : Fibre Optique FTTO (Fiber To The Office)",
            f"Débit garanti : {fibre.get('debit_mbps', '')} Mbps symétrique",
            f"Engagement contractuel : {fibre.get('engagement_mois', '')} mois",
            f"Distance raccordement : {fibre.get('distance_metres', '')} mètres",
            f"Coût d'installation : {fibre.get('cout_initial_total', 0):,.0f} TND (one-time)",
        ]
        for d in details:
            p = doc.add_paragraph(style="List Bullet")
            _run(p, d, size=9, color=DARK)
        doc.add_paragraph()

    # ── Description fonctionnelle Microsoft ───────────────────────
    ms = data.get("microsoft")
    if ms:
        _h1(doc, "SOLUTION MICROSOFT 365 — DESCRIPTION FONCTIONNELLE")
        tm = doc.add_table(rows=2, cols=4)
        tm.style = "Table Grid"
        for j, (h, v) in enumerate(zip(
            ["Plan", "Licences", "Prix unitaire/an (TND)", "Prix total/an (TND)"],
            [
                ms.get("nom_produit", "—"),
                str(ms.get("nombre_licences", "")),
                f"{ms.get('prix_unitaire_tnd', 0):,.2f}",
                f"{ms.get('prix_annuel', 0):,.0f}",
            ]
        )):
            _cell(tm, 0, j, h, bold=True, bg="FF7900", color=BLANC)
            _cell(tm, 1, j, v)
        doc.add_paragraph()

        _h2(doc, "Services inclus")
        justification = ms.get("justification", "Plan sélectionné selon les besoins du client.")
        p = doc.add_paragraph()
        _run(p, justification, size=9, italic=True, color=DARK)
        doc.add_paragraph()

    # ── Récapitulatif financier (client) ─────────────────────────
    _h1(doc, "RÉCAPITULATIF FINANCIER")
    prix_annuel  = data.get("prix_total_annuel", 0)
    prix_mensuel = round(prix_annuel / 12, 0)
    tr = doc.add_table(rows=2, cols=2)
    tr.style = "Table Grid"
    for j, (h, v) in enumerate(zip(
        ["Prix total annuel (TND)", "Équivalent mensuel (TND)"],
        [f"{prix_annuel:,.0f}", f"{prix_mensuel:,.0f}"]
    )):
        _cell(tr, 0, j, h, bold=True, bg="FF7900", color=BLANC)
        _cell(tr, 1, j, v)
    doc.add_paragraph()

    # ── Pitch + arguments ─────────────────────────────────────────
    _h1(doc, "ARGUMENTAIRE COMMERCIAL")
    pitch = data.get("pitch", "")
    if pitch:
        p = doc.add_paragraph()
        _run(p, pitch, size=10, italic=True, color=DARK)

    args = data.get("arguments", [])
    if args:
        p = doc.add_paragraph()
        _run(p, "Arguments de négociation :", size=10, bold=True, color=DARK)
        for i, arg in enumerate(args, 1):
            pa = doc.add_paragraph(style="List Number")
            _run(pa, arg, size=10)
    doc.add_paragraph()

    # ── Pied de page ──────────────────────────────────────────────
    _hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Orange Business Tunisie  |  Document confidentiel", size=8, italic=True, color=DARK)

    doc.save(chemin)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _run(para, texte, size=11, bold=False, italic=False, color=None):
    r = para.add_run(texte)
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color

def _h1(doc, texte):
    p = doc.add_paragraph()
    _run(p, texte, size=12, bold=True, color=ORANGE)

def _h2(doc, texte):
    p = doc.add_paragraph()
    _run(p, texte, size=10, bold=True, color=DARK)

def _hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "FF7900")
    pBdr.append(bot)
    pPr.append(pBdr)

def _cell(table, row, col, texte, bold=False, color=None, bg=None, row_idx=None):
    cell = table.cell(row, col)
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r    = p.add_run(str(texte))
    r.font.size = Pt(9)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor.from_string(color)
    # Lignes de données : alternance SAUMON / BLANC
    if bg is None and row_idx is not None:
        bg = "FFCCAA" if row_idx % 2 == 0 else "FFFFFF"
    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  str(bg).replace("#", ""))
        tcPr.append(shd)